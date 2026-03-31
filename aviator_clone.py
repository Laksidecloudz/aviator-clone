from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, data):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
        set_cell_shading(cell, '2E7D32')

    for row_data in data:
        row = table.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = str(text)
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    return table

def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        run = p.add_run('✓ ')
        run.font.color.rgb = RGBColor(46, 125, 50)
        run.bold = True
        p.add_run(item)

def create_report():
    doc = Document()

    # Title
    title = doc.add_heading('Development Report', level=0)
    title.alignment = 1  # Left align

    subtitle = doc.add_paragraph()
    run = subtitle.add_run('Aviator Crash Game Clone - WorldStar Betting')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    meta = doc.add_paragraph()
    run = meta.add_run('Date: March 30, 2026\nDirectory: C:\\Users\\AkayG\\Documents\\aviator-clone\nLast Updated: 2026-03-30 02:41:05+02:00')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph('-' * 80)

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'Complete UI/UX overhaul to align with Spribe Aviator reference design, '
        'including crash animation restoration, multiplier formula optimization, '
        'settings menu expansion, and responsible gaming compliance via age verification modal.'
    )

    points = [
        'Full UI/UX alignment with Spribe Aviator reference design',
        'Bug fix for multiplier sluggishness via accelerating formula',
        'Expanded settings menu with new modals and features',
        'Age verification modal for responsible gaming compliance',
        '12 files modified, 1 new file created'
    ]
    for point in points:
        doc.add_paragraph(point, style='List Bullet')

    # Section 1: UI/UX Changes
    doc.add_heading('1. UI/UX Changes (Spribe Reference Alignment)', level=1)

    doc.add_heading('1.1 Crash Animation', level=2)
    doc.add_paragraph('File: src/renderer/CanvasRenderer.js')
    add_styled_table(doc, ['Change', 'Description'], [
        ['Fly-away animation', 'Plane zooms off and fades over 1.2s'],
        ['Screen shake', 'Shake effect on crash (intensity: 12, duration: 400ms)'],
        ['Particle burst', '40 particles spawned at crash point'],
        ['Flash overlay', 'Red flash effect (alpha: 0.15)']
    ])
    doc.add_paragraph()

    doc.add_heading('1.2 Multiplier Display', level=2)
    doc.add_paragraph('Files: src/components/MultiplierDisplay.jsx, src/index.css')
    add_styled_table(doc, ['Change', 'Before', 'After'], [
        ['Font size', '80px fixed', '18vmin (responsive)'],
        ['"FLEW AWAY!"', 'Animated', 'Static white (#FFFFFF)'],
        ['Crash color', 'Red with pulse', 'Bright red (#FF0000), no animation'],
        ['Betting countdown', 'Shown in multiplier area', 'Removed']
    ])
    doc.add_paragraph()

    doc.add_heading('1.3 Color Updates', level=2)
    doc.add_paragraph('File: src/index.css')
    add_styled_table(doc, ['Element', 'Before', 'After'], [
        ['Green', '#00c945', '#29A309'],
        ['Red', '#ff3366', '#D32F2F'],
        ['Red dim', 'rgba(255, 51, 102, 0.15)', 'rgba(211, 47, 47, 0.15)']
    ])
    doc.add_paragraph()

    doc.add_heading('1.4 History Pills Color Coding', level=2)
    add_styled_table(doc, ['Multiplier Range', 'Before', 'After'], [
        ['< 2x', 'Red', 'Red'],
        ['2x - 4.99x', 'Yellow', 'Green'],
        ['5x+', 'Green', 'Pink']
    ])
    doc.add_paragraph()

    doc.add_heading('1.5 Background Color States', level=2)
    add_styled_table(doc, ['Game State', 'Color'], [
        ['LOBBY/BETTING', '#3a1a1a (Red tint)'],
        ['1x', '#1a3a6a (Blue)'],
        ['2x-5x', 'Blue to Purple gradient'],
        ['5x-10x', '#4a2a6a (Purple)'],
        ['10x+', '#6a2a5a (Pink/Magenta)'],
        ['CRASHED', 'Dark purple overlay']
    ])
    doc.add_paragraph()

    # Section 2: Bug Fix
    doc.add_heading('2. Bug Fix: Multiplier Sluggishness', level=1)
    doc.add_paragraph('File: src/game/CrashAlgorithm.js')

    doc.add_heading('Problem', level=2)
    doc.add_paragraph(
        'Multiplier felt sluggish at higher values because the formula decelerated '
        '(started fast, slowed down as it approached the crash point).'
    )

    doc.add_heading('Solution', level=2)
    doc.add_paragraph('Changed from decelerating to accelerating formula.')

    doc.add_heading('Old Formula (Decelerating)', level=3)
    doc.add_paragraph(
        'export function getMultiplierAtTime(flightTimeMs, crashMultiplier) {\n'
        '  if (crashMultiplier <= 1.00) return 1.00;\n'
        '  const t = flightTimeMs / 1000;\n'
        '  const k = computeRateConstant(crashMultiplier);\n'
        '  return 1 + (crashMultiplier - 1) * (1 - Math.exp(-k * t));\n'
        '}'
    )

    doc.add_heading('New Formula (Accelerating)', level=3)
    doc.add_paragraph(
        'export function getMultiplierAtTime(flightTimeMs, crashMultiplier) {\n'
        '  if (crashMultiplier <= 1.00) return 1.00;\n'
        '  const t = flightTimeMs / 1000;\n'
        '  const k = computeRateConstant(crashMultiplier);\n'
        '  const normFactor = Math.exp(k * 8) - 1;\n'
        '  return 1 + (crashMultiplier - 1) * (Math.exp(k * t) - 1) / normFactor;\n'
        '}'
    )

    doc.add_heading('Impact', level=2)
    for item in ['Multiplier now accelerates throughout flight', 'Creates tension as it speeds toward crash point', 'Matches real Spribe Aviator feel']:
        doc.add_paragraph(item, style='List Bullet')

    # Section 3: Settings Menu
    doc.add_heading('3. Settings Menu Expansion', level=1)
    doc.add_paragraph('File: src/components/SettingsMenu.jsx (full rewrite)')

    doc.add_heading('New Modals Added', level=2)
    add_styled_table(doc, ['Menu Modal', 'Description'], [
        ['My Bet History', '10 mock bets, stats (total wagered/won), scrollable list'],
        ['Free Bets & Rain', '4 mock free bets with claim buttons, expiration timers'],
        ['Game Limits', 'Bet limits ($1-$100), max win ($10,000), auto cashout limits'],
        ['How To Play', '4-step visual guide with tip box'],
        ['Game Rules', 'Basic Rules, Multipliers, Payouts, Auto Features sections']
    ])
    doc.add_paragraph()

    doc.add_heading('Enhancements', level=2)
    for item in ['Player avatar shows initial letter', 'Free Bets badge shows count of claimable bets', 'Modal animations (slide up, fade backdrop)', 'Footer shows version number']:
        doc.add_paragraph(item, style='List Bullet')

    # Section 4: Welcome Modal
    doc.add_heading('4. Welcome Modal (Age Verification)', level=1)
    doc.add_paragraph('File: src/components/WelcomeModal.jsx [NEW]')

    doc.add_heading('Features', level=2)
    add_styled_table(doc, ['Feature', 'Description'], [
        ['Gambling warning', 'Red box with warning text'],
        ['Responsible gaming', 'Green box with tips (set limits, dont chase losses)'],
        ['Help link', 'BeGambleAware.org'],
        ['Age checkbox', '18+ verification with custom styled checkbox'],
        ['Enter button', 'Disabled until checkbox checked'],
        ['Persistence', 'localStorage key: aviator_age_confirmed']
    ])
    doc.add_paragraph()

    doc.add_heading('User Flow', level=2)
    for i, step in enumerate([
        'User visits site',
        'Welcome modal appears (blocks game access)',
        'User reads warnings',
        'User checks "I am 18 years or older"',
        'User clicks "Enter Game"',
        'Modal closes, localStorage set',
        'Modal wont show again on return visits'
    ], 1):
        doc.add_paragraph(f'{i}. {step}')

    # Section 5: Files Modified
    doc.add_heading('5. Files Modified', level=1)
    add_styled_table(doc, ['#', 'File', 'Type', 'Changes'], [
        ['1', 'src/App.jsx', 'Modified', 'Added WelcomeModal integration'],
        ['2', 'src/index.css', 'Modified', 'Color updates, modal styles (+200 lines)'],
        ['3', 'src/components/BetPanel.jsx', 'Modified', 'Quick amounts, button formats'],
        ['4', 'src/components/BetTable.jsx', 'Modified', 'Player count, avatars, headers'],
        ['5', 'src/components/GameHistory.jsx', 'Modified', 'Color coding (red/green/pink)'],
        ['6', 'src/components/MultiplierDisplay.jsx', 'Modified', 'Removed betting countdown'],
        ['7', 'src/components/BalanceDisplay.jsx', 'Modified', 'USD format with toLocaleString'],
        ['8', 'src/components/SponsorBanner.jsx', 'Modified', 'aviator logo with SVG propeller'],
        ['9', 'src/components/SettingsMenu.jsx', 'Modified', 'Full menu expansion with modals'],
        ['10', 'src/components/WelcomeModal.jsx', 'NEW', 'Age verification modal'],
        ['11', 'src/renderer/CanvasRenderer.js', 'Modified', 'Fly-away, shake, particles'],
        ['12', 'src/renderer/BackgroundRenderer.js', 'Modified', 'Dynamic gradients'],
        ['13', 'src/game/CrashAlgorithm.js', 'Modified', 'Accelerating multiplier formula']
    ])
    doc.add_paragraph()
    p = doc.add_paragraph('Total: 12 modified, 1 new file')
    p.runs[0].italic = True

    # Section 6: Build Status
    doc.add_heading('6. Build & Lint Status', level=1)

    doc.add_heading('Build', level=2)
    p = doc.add_paragraph()
    run = p.add_run('built in 586ms\ndist/index.html          0.66 kB\ndist/assets/index.css    25.59 kB\ndist/assets/index.js     245.10 kB')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_heading('Lint', level=2)
    doc.add_paragraph('Passing (only legacy files in aviator-betting-game-clone-master/ show warnings)')

    # Section 7: Testing Checklist
    doc.add_heading('7. Testing Checklist', level=1)

    doc.add_heading('Core Gameplay', level=2)
    add_checklist(doc, [
        'Fly-away animation on crash',
        'Screen shake and particles on crash',
        '"FLEW AWAY!" white text, red multiplier',
        'Background color transitions',
        'Accelerating multiplier formula'
    ])

    doc.add_heading('UI Components', level=2)
    add_checklist(doc, [
        'Bet panel formats (USD, quick amounts 1/2/5/10)',
        'Left sidebar (player count, avatars, headers)',
        'History pills colors (red/green/pink)',
        'Balance format ("30,000.88 USD")',
        'Sponsor banner ("aviator" logo)'
    ])

    doc.add_heading('Menus', level=2)
    add_checklist(doc, [
        'Settings menu opens/closes',
        'Free Bets modal with claim buttons',
        'Bet History modal with stats',
        'Game Limits modal',
        'How To Play modal',
        'Game Rules modal',
        'Welcome modal with age verification',
        'localStorage persistence'
    ])

    # Section 8: Technical Notes
    doc.add_heading('8. Technical Notes', level=1)

    doc.add_heading('State Management', level=2)
    doc.add_paragraph('showWelcome state initialized from localStorage', style='List Bullet')
    doc.add_paragraph('aviator_age_confirmed key used for persistence', style='List Bullet')

    doc.add_heading('Performance', level=2)
    doc.add_paragraph('Background renderer calculates colors per frame based on multiplier', style='List Bullet')
    doc.add_paragraph('Accelerating formula uses Math.exp for smooth curves', style='List Bullet')

    doc.add_heading('Browser Compatibility', level=2)
    doc.add_paragraph('Uses localStorage for age verification persistence', style='List Bullet')
    doc.add_paragraph('toLocaleString for number formatting', style='List Bullet')

    # Footer
    doc.add_paragraph('-' * 80)
    footer = doc.add_paragraph()
    run = footer.add_run('Report Generated: 2026-03-30T02:41:05+02:00\nAviator Crash Game Clone - WorldStar Betting')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)
    run.italic = True

    # Save
    doc.save('Development_Report_Aviator_Clone.docx')
    print('Done! Report saved as Development_Report_Aviator_Clone.docx')

if __name__ == '__main__':
    create_report()
